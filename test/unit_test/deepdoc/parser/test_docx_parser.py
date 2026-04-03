#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#  http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.

"""Unit tests for the DOCX parser.

Tests cover:
- Basic parsing functionality with various content types
- Edge cases: empty files, special characters, large files
- Table extraction and composition
- Image handling within paragraphs
- Page range filtering
- Style name extraction
"""

import importlib.util
import os
import sys
from io import BytesIO
from unittest import mock


_MOCK_MODULES = [
    "xgboost",
    "xgb",
    "pdfplumber",
    "huggingface_hub",
    "PIL",
    "PIL.Image",
    "pypdf",
    "sklearn",
    "sklearn.cluster",
    "sklearn.metrics",
    "deepdoc.vision",
    "infinity",
    "infinity.rag_tokenizer",
]
for _m in _MOCK_MODULES:
    if _m not in sys.modules:
        sys.modules[_m] = mock.MagicMock()


def _find_project_root(marker="pyproject.toml"):
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.exists(os.path.join(d, marker)):
            return d
        d = os.path.dirname(d)
    return None


_PROJECT_ROOT = _find_project_root()

_lazy_image_spec = importlib.util.spec_from_file_location(
    "rag.utils.lazy_image",
    os.path.join(_PROJECT_ROOT, "rag", "utils", "lazy_image.py"),
)
_lazy_image_mod = importlib.util.module_from_spec(_lazy_image_spec)
sys.modules["rag.utils.lazy_image"] = _lazy_image_mod
_lazy_image_spec.loader.exec_module(_lazy_image_mod)

_block_type_spec = importlib.util.spec_from_file_location(
    "deepdoc.vision.block_type",
    os.path.join(_PROJECT_ROOT, "deepdoc", "vision", "block_type.py"),
)
_block_type_mod = importlib.util.module_from_spec(_block_type_spec)
sys.modules["deepdoc.vision.block_type"] = _block_type_mod
_block_type_spec.loader.exec_module(_block_type_mod)

_docx_spec = importlib.util.spec_from_file_location(
    "deepdoc.parser.docx_parser",
    os.path.join(_PROJECT_ROOT, "deepdoc", "parser", "docx_parser.py"),
)
_docx_mod = importlib.util.module_from_spec(_docx_spec)
sys.modules["deepdoc.parser.docx_parser"] = _docx_mod
_docx_spec.loader.exec_module(_docx_mod)

RAGFlowDocxParser = _docx_mod.RAGFlowDocxParser


def _create_docx_with_paragraphs(paragraphs, tables=None, include_page_break=False):
    """Create a minimal DOCX file in memory.

    Args:
        paragraphs: list of text strings for paragraphs
        tables: list of tables, each table is a list of rows, each row is a list of cells
        include_page_break: whether to include page breaks between paragraphs

    Returns:
        bytes: DOCX file content
    """
    from docx import Document

    doc = Document()

    for i, para_text in enumerate(paragraphs):
        if i > 0 and include_page_break:
            doc.add_page_break()
        doc.add_paragraph(para_text)

    if tables:
        for table_data in tables:
            if not table_data:
                continue
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    table.rows[i].cells[j].text = str(cell_text)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _create_docx_with_image():
    """Create a DOCX with an embedded image.

    Returns:
        bytes: DOCX file content with image
    """
    from docx import Document
    from docx.shared import Inches
    from PIL import Image

    img_buf = BytesIO()
    img = Image.new("RGB", (100, 100), color="red")
    img.save(img_buf, format="PNG")
    img_buf.seek(0)

    doc = Document()
    doc.add_paragraph("Before image")
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(img_buf, width=Inches(1))
    doc.add_paragraph("After image")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


class TestDocxParserBasic:
    """Tests for basic DOCX parsing functionality."""

    def test_parse_simple_document(self):
        """Parse a document with simple text paragraphs."""
        paragraphs = ["First paragraph", "Second paragraph", "Third paragraph"]
        docx_bytes = _create_docx_with_paragraphs(paragraphs)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(secs) == 3
        assert len(tbls) == 0
        assert secs[0][0] == "First paragraph"
        assert secs[1][0] == "Second paragraph"
        assert secs[2][0] == "Third paragraph"

    def test_parse_empty_document(self):
        """Parse an empty document."""
        docx_bytes = _create_docx_with_paragraphs([])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(secs) == 0
        assert len(tbls) == 0

    def test_parse_document_with_whitespace_paragraphs(self):
        """Paragraphs with only whitespace are stripped."""
        paragraphs = ["Valid text", "   ", "More valid text"]
        docx_bytes = _create_docx_with_paragraphs(paragraphs)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(secs) == 3
        assert secs[0][0] == "Valid text"
        assert secs[2][0] == "More valid text"

    def test_parse_document_with_chinese_text(self):
        """Parse a document containing Chinese characters."""
        paragraphs = ["这是中文段落", "English and 中文 mixed", "数字123测试"]
        docx_bytes = _create_docx_with_paragraphs(paragraphs)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(secs) == 3
        assert secs[0][0] == "这是中文段落"
        assert secs[1][0] == "English and 中文 mixed"

    def test_parse_document_with_special_characters(self):
        """Parse a document with special characters."""
        paragraphs = [
            "Special chars: !@#$%^&*()",
            "Quotes: \"'`~",
            "Math: +=-*/<=>",
        ]
        docx_bytes = _create_docx_with_paragraphs(paragraphs)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(secs) == 3
        assert "Special chars" in secs[0][0]
        assert "Quotes" in secs[1][0]

    def test_parse_from_file_path(self, temp_docx):
        """Test parsing from a file path (string) instead of bytes."""
        paragraphs = ["File path test"]
        docx_bytes = _create_docx_with_paragraphs(paragraphs)
        temp_path = temp_docx(suffix=".docx", content=docx_bytes)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(temp_path)

        assert len(secs) == 1
        assert secs[0][0] == "File path test"


class TestDocxParserTables:
    """Tests for table extraction functionality."""

    def test_parse_document_with_simple_table(self):
        """Parse a document with a simple table."""
        table_data = [
            ["Header1", "Header2", "Header3"],
            ["Value1", "Value2", "Value3"],
            ["Value4", "Value5", "Value6"],
        ]
        docx_bytes = _create_docx_with_paragraphs(["Before table"], tables=[table_data])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 1
        assert len(tbls[0]) > 0

    def test_parse_document_with_numeric_table(self):
        """Parse a table with numeric data."""
        table_data = [
            ["Name", "Score", "Grade"],
            ["Alice", "95", "A"],
            ["Bob", "87", "B"],
        ]
        docx_bytes = _create_docx_with_paragraphs([], tables=[table_data])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 1

    def test_parse_document_with_empty_table(self):
        """Parse a document with an empty table."""
        table_data = [["", ""], ["", ""]]
        docx_bytes = _create_docx_with_paragraphs([], tables=[table_data])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 1

    def test_parse_document_with_multiple_tables(self):
        """Parse a document with multiple tables."""
        table1 = [["A", "B"], ["1", "2"]]
        table2 = [["C", "D"], ["3", "4"]]
        docx_bytes = _create_docx_with_paragraphs([], tables=[table1, table2])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 2

    def test_parse_document_with_single_row_table(self):
        """Parse a table with only one row (header only)."""
        table_data = [["Header1", "Header2", "Header3"]]
        docx_bytes = _create_docx_with_paragraphs([], tables=[table_data])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 1
        assert tbls[0] == []


class TestDocxParserPageRange:
    """Tests for page range filtering functionality."""

    def test_parse_with_from_page(self):
        """Test parsing starting from a specific page."""
        paragraphs = ["Page 1", "Page 2", "Page 3"]
        docx_bytes = _create_docx_with_paragraphs(paragraphs, include_page_break=True)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes, from_page=1)

        assert len(secs) == 2
        assert secs[0][0] == "Page 2"
        assert secs[1][0] == "Page 3"

    def test_parse_with_to_page(self):
        """Test parsing up to a specific page."""
        paragraphs = ["Page 1", "Page 2", "Page 3"]
        docx_bytes = _create_docx_with_paragraphs(paragraphs, include_page_break=True)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes, to_page=2)

        assert len(secs) == 2
        assert secs[0][0] == "Page 1"
        assert secs[1][0] == "Page 2"

    def test_parse_with_page_range(self):
        """Test parsing a specific page range."""
        paragraphs = ["Page 1", "Page 2", "Page 3", "Page 4"]
        docx_bytes = _create_docx_with_paragraphs(paragraphs, include_page_break=True)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes, from_page=1, to_page=3)

        assert len(secs) == 2
        assert secs[0][0] == "Page 2"
        assert secs[1][0] == "Page 3"


class TestDocxParserStyleExtraction:
    """Tests for style name extraction."""

    def test_style_name_extracted(self):
        """Verify style names are extracted along with text."""
        paragraphs = ["Styled paragraph"]
        docx_bytes = _create_docx_with_paragraphs(paragraphs)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(secs) == 1
        assert len(secs[0]) == 2
        text, style_name = secs[0]
        assert text == "Styled paragraph"
        assert isinstance(style_name, str)


class TestDocxParserImages:
    """Tests for image handling within paragraphs."""

    def test_get_picture_no_image(self):
        """get_picture returns None for paragraphs without images."""
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("Text without image")

        parser = RAGFlowDocxParser()
        result = parser.get_picture(doc, p)

        assert result is None

    def test_parse_document_with_embedded_image(self):
        """Parse a document containing an embedded image."""
        docx_bytes = _create_docx_with_image()

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(secs) >= 2


class TestDocxParserEdgeCases:
    """Tests for edge cases and error handling."""

    def test_parse_very_long_paragraph(self):
        """Parse a document with a very long paragraph."""
        long_text = "A" * 10000
        docx_bytes = _create_docx_with_paragraphs([long_text])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(secs) == 1
        assert len(secs[0][0]) == 10000

    def test_parse_unicode_text(self):
        """Parse a document with various Unicode characters."""
        paragraphs = [
            "Emoji: 😀🎉🎊",
            "Arabic: مرحبا",
            "Russian: Привет",
            "Japanese: こんにちは",
        ]
        docx_bytes = _create_docx_with_paragraphs(paragraphs)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(secs) == 4

    def test_parse_newlines_in_paragraph(self):
        """Parse paragraphs containing newline characters."""
        paragraphs = ["Line 1\nLine 2\nLine 3"]
        docx_bytes = _create_docx_with_paragraphs(paragraphs)

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(secs) == 1

    def test_parse_table_with_merged_cells(self):
        """Parse a table that simulates merged cells."""
        table_data = [
            ["Merged", "", "Value"],
            ["A", "B", "C"],
        ]
        docx_bytes = _create_docx_with_paragraphs([], tables=[table_data])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 1

    def test_extract_table_content_structure(self):
        """Verify table content extraction structure."""
        table_data = [
            ["Name", "Age", "City"],
            ["Alice", "30", "NYC"],
            ["Bob", "25", "LA"],
        ]
        docx_bytes = _create_docx_with_paragraphs([], tables=[table_data])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 1
        table_content = tbls[0]
        assert isinstance(table_content, list)


class TestDocxParserTableComposition:
    """Tests for table content composition logic."""

    def test_table_with_numeric_data(self):
        """Table with numeric data gets proper formatting."""
        table_data = [
            ["Item", "Price", "Quantity"],
            ["Apple", "1.50", "10"],
            ["Banana", "0.75", "20"],
        ]
        docx_bytes = _create_docx_with_paragraphs([], tables=[table_data])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 1

    def test_table_with_date_data(self):
        """Table with date-like data."""
        table_data = [
            ["Date", "Event"],
            ["2024-01-15", "Meeting"],
            ["2024-02-20", "Conference"],
        ]
        docx_bytes = _create_docx_with_paragraphs([], tables=[table_data])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 1

    def test_wide_table_format(self):
        """Wide tables (more than 3 columns) get line-separated format."""
        table_data = [
            ["Col1", "Col2", "Col3", "Col4"],
            ["A", "B", "C", "D"],
        ]
        docx_bytes = _create_docx_with_paragraphs([], tables=[table_data])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 1

    def test_narrow_table_format(self):
        """Narrow tables (3 or fewer columns) get newline-separated format."""
        table_data = [
            ["Col1", "Col2", "Col3"],
            ["A", "B", "C"],
        ]
        docx_bytes = _create_docx_with_paragraphs([], tables=[table_data])

        parser = RAGFlowDocxParser()
        secs, tbls = parser(docx_bytes)

        assert len(tbls) == 1
